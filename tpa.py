from openpyxl import load_workbook
from docx import Document

words_fs = ['{S1}', '{S2}', '{S3}', '{S4}', '{S5}', '{S6}', '{S7}']

def saveData(grsName, grsData):
    document = Document('form.docx')    
    for i in range(len(words_fs)):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if words_fs[i] in cell.text:
                        currentAligment = cell.paragraphs[0].alignment
                        currentFont = cell.paragraphs[0].runs[0].font.name                    
                        cell.text = grsData[i]
                        cell.paragraphs[0].alignment = currentAligment
                        cell.paragraphs[0].runs[0].font.name = currentFont
    document.save(grsName + ', ' + grsData[6] + '.docx')

def useData(first_line, last_line):
    wb = load_workbook('tpa.xlsx')
    ws = wb.active    

    # i = current line
    for i in range(first_line, last_line+1):
        # prepairig datalist for form
        wfc = []
        # name of GRS
        d = ws['A' + str(i)]
        current_grs = d.value        
        # type of tpa
        d = ws['F' + str(i)]
        wfc.append(d.value)
        # marka tpa
        d = ws['G' + str(i)]
        wfc.append(d.value)
        # Dy tpa
        d = ws['H' + str(i)]
        wfc.append(str(d.value))
        # Py tpa
        d = ws['I' + str(i)]
        currentText = str(d.value / 10)
        currentText = currentText.replace('.', ',')
        wfc.append(currentText)
        # Working gas    
        wfc.append('Природный газ')
        # Date of build
        d = ws['J' + str(i)]
        dt = d.value
        d = dt.strftime("%d.%m.%Y")
        wfc.append(d)
        # Place and number
        d = ws['C' + str(i)]
        d1 = ws['D' + str(i)]
        wfc.append(str(d.value + ', №' + d1.value))               
        saveData(current_grs, wfc)
        print(str(i - first_line + 1) + '/' + str(last_line - first_line + 1))
    
fLine = input("Введите номер начальной строки таблицы: ")
lLine = input("Введите номер последней строки таблицы: ")
useData(int(fLine), int(lLine))
