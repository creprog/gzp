from openpyxl import load_workbook
from docx import Document

words_fs = ['{S1}', '{S2}', '{S3}', '{S4}', '{S5}', '{S6}', '{S7}']

def saveData(grsData):
    document = Document('form.docx')    
    for i in range(len(words_fs)):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if words_fs[i] in cell.text:
                        currentAligment = cell.paragraphs[0].alignment
                        currentFont = cell.paragraphs[0].runs[0].font.name                    
                        cell.text = grsData[i]
                        cell.paragraphs[0].alignment = currentAligment
                        cell.paragraphs[0].runs[0].font.name = currentFont
    document.save(grsData[6] + '.docx')

def useData(first_line, last_line):
    wb = load_workbook('ppk.xlsx')
    ws = wb.active    

    # i = current line
    for i in range(first_line, last_line+1):
        try:
            # prepairig datalist for form
            wfc = []
            # name of GRS
            d = ws['A' + str(i)]
            current_grs = d.value        
            # type of tpa
            tpatype = 'Клапан предохранительный'
            wfc.append(tpatype)
            # marka tpa
            try:
                d = ws['F' + str(i)]
                wfc.append(d.value)
            except:
                wfc.append(' ')
            # Dy tpa
            d = ws['G' + str(i)]
            wfc.append(str(d.value))
            # Py tpa
            d = ws['H' + str(i)]
            currentText = str(d.value / 10)
            currentText = currentText.replace('.', ',')
            wfc.append(currentText)
            # Working gas    
            wfc.append('Природный газ')
            # Date of build
            d = ws['L' + str(i)]
            dt = str(d.value) + 'г'
            #d = dt.strftime("%d.%m.%Y")
            wfc.append(dt)
            # Place and number
            d = ws['C' + str(i)]
            d1 = ws['D' + str(i)]
            full_line = 'ГРС ' + current_grs + ', ' + str(d.value).capitalize() + ', №' + str(d1.value)
            wfc.append(full_line)
            saveData(wfc)
            print(str(i) + '/' + str(last_line))
        except:
            print('Не сделан паспорт для строки: ' + str(i))
    
fLine = input("Введите номер начальной строки таблицы: ")
lLine = input("Введите номер последней строки таблицы: ")
useData(int(fLine), int(lLine))
