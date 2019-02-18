from openpyxl import load_workbook
from docx import Document

words_fs = ['{S1}', '{S2}', '{S3}', '{S4}', '{S5}', '{S6}', '{S7}']

py_dict = {
    '6.276256': '6,4',
    '5.88399': '6,0',
    '15.69064': '16,0',
    '9.80665': '10,0',
    '1.569064': '1,6',
    '1.176798': '1,2',
    '7.84532': '8,0',
    '0.588399': '0,6',
    '4.903325': '5,0',
    '0.6864655': '1,0',
    '7.3549875': '7,5',
}

def clear_name(grs_name):
    for i in range(len(grs_name)):
        if grs_name[i] == '(':
            return grs_name[:i-1]

def printlst(grsname, somelist):
    print(grsname)
    for i in somelist:
        print(i)

def saveData(grsName, grsData):
    document = Document('form.docx')    
    for i in range(len(words_fs)):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if words_fs[i] in cell.text:
                        currentAligment = cell.paragraphs[0].alignment
                        currentFont = cell.paragraphs[0].runs[0].font.name
                        cell.text = grsData[i]
                        cell.paragraphs[0].alignment = currentAligment
                        cell.paragraphs[0].runs[0].font.name = currentFont
    document.save(grsName + ', ' + grsData[6] + '.docx')

def readFile(first_line, last_line):
    wb = load_workbook('rd.xlsx')
    ws = wb.active
    current_grs = ''
    # i = current line
    for i in range(first_line, last_line+1):
        # prepairig datalist for form
        wfc = []        
        rd_number = 1
        # name of GRS
        try:
            d = ws['C' + str(i)]            
            current_grs = clear_name(d.value)            
        except:
            current_grs = current_grs
        
        # type of rd
        d = ws['D' + str(i)]
        wfc.append(d.value)

        # marka rd
        d = ws['E' + str(i)]
        if d.value != None:
            wfc.append(d.value)
            rd_number = 1
        else:
            d = ws['F' + str(i)]
            wfc.append(d.value)
            rd_number = 2
        # Dy rd
        d = ws['G' + str(i)]
        wfc.append(str(d.value))
        # Py tpa
        d = ws['H' + str(i)]        
        wfc.append(py_dict[str(d.value)])
        # Working gas    
        wfc.append('Природный газ')
        # Date of build
        d = ws['M' + str(i)]
        dt = d.value
        d = dt.strftime("%d.%m.%Y")
        wfc.append(d)
        # Place and number
        d = ws['A' + str(i)]
        rd_line = ws['I' + str(i)]
        rd_line = rd_line.value
        d1 = 'РД4.' + str(rd_line) + '-' + str(rd_number)
        wfc.append(str(d.value + ', №' + d1))
        #printlst(current_grs, wfc)
        saveData(current_grs, wfc)
        print(str(i - first_line + 1) + '/' + str(last_line - first_line + 1))

fLine = input("Введите номер начальной строки таблицы: ")
lLine = input("Введите номер последней строки таблицы: ")
readFile(int(fLine), int(lLine))
